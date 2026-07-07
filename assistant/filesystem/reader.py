"""File readers for supported document formats."""

from __future__ import annotations

import csv
import json
from pathlib import Path
from typing import Any
from xml.etree import ElementTree

import yaml

from assistant.filesystem.validators import FilesystemValidators


class FilesystemReader:
    """Read supported file formats into structured data."""

    TEXT_EXTENSIONS = {".txt", ".md", ".py", ".java", ".cpp", ".c", ".html", ".css", ".js"}

    def __init__(self, maximum_file_size: int) -> None:
        self.maximum_file_size = maximum_file_size

    def read(self, path: Path, preview_rows: int = 20) -> dict[str, Any]:
        """Read a file and return format-specific content."""
        resolved = path.expanduser().resolve()
        FilesystemValidators.readable_file(resolved, self.maximum_file_size)
        extension = resolved.suffix.lower()
        if extension in self.TEXT_EXTENSIONS:
            return self._read_text(resolved)
        if extension == ".pdf":
            return self._read_pdf(resolved)
        if extension == ".docx":
            return self._read_docx(resolved)
        if extension == ".csv":
            return self._read_csv(resolved, preview_rows)
        if extension == ".json":
            return {"path": str(resolved), "content": json.loads(resolved.read_text(encoding="utf-8"))}
        if extension in {".yaml", ".yml"}:
            return {"path": str(resolved), "content": yaml.safe_load(resolved.read_text(encoding="utf-8"))}
        if extension == ".xml":
            root = ElementTree.parse(resolved).getroot()
            return {"path": str(resolved), "root_tag": root.tag, "text": "".join(root.itertext())}
        raise ValueError(f"Unsupported file format: {extension or 'unknown'}")

    def _read_text(self, path: Path) -> dict[str, Any]:
        """Read UTF-8-ish text with replacement for invalid bytes."""
        return {"path": str(path), "content": path.read_text(encoding="utf-8", errors="replace")}

    def _read_pdf(self, path: Path) -> dict[str, Any]:
        """Read PDF text using PyMuPDF."""
        try:
            import fitz
        except ImportError as exc:
            raise ValueError("PDF support requires PyMuPDF") from exc
        with fitz.open(path) as document:
            text = "\n".join(page.get_text() for page in document)
            return {
                "path": str(path),
                "content": text,
                "page_count": document.page_count,
                "metadata": dict(document.metadata or {}),
            }

    def _read_docx(self, path: Path) -> dict[str, Any]:
        """Read DOCX paragraphs and tables using python-docx."""
        try:
            import docx
        except ImportError as exc:
            raise ValueError("DOCX support requires python-docx") from exc
        document = docx.Document(path)
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
        tables = [
            [[cell.text for cell in row.cells] for row in table.rows]
            for table in document.tables
        ]
        return {"path": str(path), "paragraphs": paragraphs, "tables": tables}

    def _read_csv(self, path: Path, preview_rows: int) -> dict[str, Any]:
        """Read CSV headers and a bounded preview."""
        with path.open(newline="", encoding="utf-8", errors="replace") as handle:
            reader = csv.reader(handle)
            rows = list(reader)
        headers = rows[0] if rows else []
        body = rows[1:]
        return {
            "path": str(path),
            "headers": headers,
            "rows": body,
            "column_count": len(headers),
            "preview": body[: max(0, preview_rows)],
        }
